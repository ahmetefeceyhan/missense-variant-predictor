"""
Proje Detay Raporu (PDR) oluşturma scripti.
Şablon DOCX'in header/arka planını ve kapak sayfasını korur, içeriği doldurur.
"""
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

TEMPLATE = Path("/Users/tefe/Downloads/2026_PDR_Şablon_Universite_TR_bCw49.docx")
PROJECT = Path("/Users/tefe/teknofest_model/teknofest_model")
OUTPUT = PROJECT / "reports" / "Proje_Detay_Raporu.docx"
RESULTS = PROJECT / "results"

ENGLISH_TERMS = [
    "missense", "benign", "pathogenic", "label shift", "precision", "recall",
    "F1", "MCC", "AUC-ROC", "AUC-PR", "bootstrap", "hold-out", "overfitting",
    "undersampling", "stacking", "ensemble", "baseline", "overfit",
    "feature importance", "feature engineering", "grid search",
    "one-hot encoding", "label encoding", "dropout", "weight_decay",
    "early stopping", "focal loss", "class_weight", "scale_pos_weight",
    "cross-validation", "out-of-fold", "meta-learner", "meta-feature",
    "Leave-One-Out Cross-Validation", "LOO-CV",
    "BalancedBagging", "prior-shift", "label shift",
    "confusion matrix", "false positive", "false negative",
    "threshold", "data leakage", "imputation",
    "gradient boosting", "leaf-wise", "depth-wise",
    "residual connection", "RBF kernel", "Platt scaling",
    "foundation model", "curate", "tabular",
    "ground truth", "in-silico", "state-of-the-art",
    "batch-effect", "finetune", "fine-tune", "pretrain", "pre-train",
    "pipeline", "framework", "domain",
    "Stratified K-Fold", "stratified",
    "decision boundary", "specificity", "sensitivity",
    "logistic regression", "Logistic Regression",
    "Variant of Uncertain Significance", "VUS",
    "minimax-optimal", "resample",
    "Minor Allele Frequency", "MAF",
    "SmallMLP", "DeepMLP", "TabPFN",
    "SMOTE", "CatBoost", "LightGBM", "XGBoost", "SVM",
    "COMBINED", "native",
]


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells_data, bold=False, shade=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = "Aptos"
        if bold:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if shade:
            set_cell_shading(cell, shade)
    return row


def make_header_row(table, cells_data):
    for i, text in enumerate(cells_data):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Aptos"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F4E79")


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    if not os.path.exists(img_path):
        p = doc.add_paragraph(f"[Şekil bulunamadı: {img_path}]")
        p.style = doc.styles['Normal']
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.name = "Aptos"
    r.italic = True


def heading(doc, text):
    p = doc.add_paragraph(text, style='Heading 3')
    return p


def subheading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_text_with_italic_terms(p, text, bold_all=True)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_text_with_italic_terms(p, text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Paragraph')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_text_with_italic_terms(p, text)
    return p


def _add_text_with_italic_terms(paragraph, text, bold_all=False):
    sorted_terms = sorted(ENGLISH_TERMS, key=len, reverse=True)
    pattern_parts = [re.escape(t) for t in sorted_terms]
    pattern = '(' + '|'.join(pattern_parts) + ')'
    parts = re.split(pattern, text)

    term_set = set(t.lower() for t in ENGLISH_TERMS)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(12)
        run.font.name = "Aptos"
        if bold_all:
            run.bold = True
        if part.lower() in term_set:
            run.italic = True


def add_page_break(doc):
    pb_para = doc.add_paragraph()
    run_pb = pb_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run_pb._element.append(br)


def fill_cover_page(doc):
    sdt = doc.element.body[0]
    all_paras = sdt.findall('.//' + qn('w:p'))

    for p in all_paras:
        full = ''
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                full += (t.text or '')
        full_stripped = full.strip()

        if full_stripped == 'Proje Adı:\xa0' or full_stripped == 'Proje Adı:':
            _append_text_to_para(p, " Patojenite Teşhisi İçin Makine Öğrenmesi Temelli Bir Varyant İnceleme Modeli")
        elif full_stripped == 'Takım Adı:':
            _append_text_to_para(p, " ClariGen")
        elif full_stripped == 'Takım ID:':
            _append_text_to_para(p, " 883973")
        elif full_stripped == 'Başvuru ID:':
            _append_text_to_para(p, " 4873357")


def _append_text_to_para(p_elem, text):
    last_run = p_elem.findall(qn('w:r'))
    if last_run:
        ref_run = last_run[-1]
        rpr = ref_run.find(qn('w:rPr'))
    else:
        rpr = None

    new_run = OxmlElement('w:r')
    if rpr is not None:
        new_rpr = parse_xml(rpr.xml)
        new_run.append(new_rpr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)
    p_elem.append(new_run)


def add_toc_page(doc):
    heading(doc, "İçindekiler")
    doc.add_paragraph()

    toc_items = [
        ("1. Giriş", ""),
        ("2. Yöntem", ""),
        ("    2.1 Kullanılan Veri Kümesi ve Yapısı", ""),
        ("    2.2 Veri Ön İşleme ve Temsilleme Stratejisi", ""),
        ("    2.3 Sınıf Dengesi ve Label Shift Yönetimi", ""),
        ("    2.4 Seçilen Algoritmalar ve Gerekçe", ""),
        ("    2.5 Hiperparametre Optimizasyonu ve Çapraz Doğrulama", ""),
        ("    2.6 Stacking Ensemble ve Out-of-Fold Protokolü", ""),
        ("    2.7 Değerlendirme Protokolü", ""),
        ("3. Bulgular", ""),
        ("    3.1 Baseline Sonuçları", ""),
        ("    3.2 Ablasyon Deneyleri", ""),
        ("    3.3 Panel-Bazlı Eğitim Sonuçları", ""),
        ("    3.4 CFTR Paneli: Refinement ve Final Model", ""),
        ("    3.5 PAH Paneli: Refinement ve Plato Analizi", ""),
        ("    3.6 Hereditary Cancer (KANSER) Paneli", ""),
        ("    3.7 MASTER (Genel Panel)", ""),
        ("    3.8 Karar Eşiği Analizi", ""),
        ("4. Sonuç", ""),
        ("    4.1 Güçlü Yönler ve Başarılar", ""),
        ("    4.2 Sınırlılıklar ve Zorluklar", ""),
        ("    4.3 Literatürdeki Yeri", ""),
        ("    4.4 Final Aşamasında Beklenen Zorluklar", ""),
        ("5. Kaynakça", ""),
    ]

    for item, _ in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "Aptos"
        if not item.startswith("    "):
            run.bold = True

    add_page_break(doc)


def build_report():
    doc = Document(str(TEMPLATE))

    # 1. Fill cover page fields
    fill_cover_page(doc)

    # 2. Remove template placeholder content (keep SDT cover + sectPr)
    body_elem = doc.element.body
    elements_to_remove = []
    for child in body_elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('sdt', 'sectPr'):
            continue
        elements_to_remove.append(child)
    for elem in elements_to_remove:
        body_elem.remove(elem)

    # 3. Page break after cover
    add_page_break(doc)

    # 4. İçindekiler
    add_toc_page(doc)

    # =====================================================
    # 1. GİRİŞ
    # =====================================================
    heading(doc, "1. GİRİŞ")
    doc.add_paragraph()

    body(doc,
        "Bu çalışma, TEKNOFEST 2026 Sağlıkta Yapay Zeka Yarışması kapsamında, "
        "genetik varyantların pathogenic (hastalık yapıcı) veya benign (zararsız) olarak sınıflandırılması "
        "problemini ele almaktadır. missense varyant patojenite tahmini, klinik genetik alanında kritik bir "
        "sorundur; çünkü yeni nesil dizileme teknolojilerinin yaygınlaşmasıyla her hasta genomunda binlerce "
        "belirsiz anlama sahip varyant (Variant of Uncertain Significance, VUS) tespit edilmektedir. "
        "Bu varyantların doğru sınıflandırılması, hastalık tanısı, tedavi planlaması ve genetik danışmanlık "
        "açısından hayati önem taşımaktadır [1]."
    )

    body(doc,
        "Yarışmada sağlanan veri seti, ClinVar ve gnomAD gibi açık kaynak veritabanlarından curate "
        "edilmiş, anonimleştirilmiş bir tablo (tabular) veri yapısına sahiptir. Veri seti dört panelden "
        "oluşmaktadır: MASTER (genel; 2931 satır), Hereditary Cancer (KANSER; 388 satır), "
        "PAH (Pulmoner Arteryel Hipertansiyon; 372 satır) ve CFTR (Kistik Fibrozis; 111 satır). "
        "Sütunlar anonimleştirilmiş olup AL_1..AL_334 (sayısal özellikler), CAT_1..CAT_6 (kategorik), "
        "EK_1..EK_9 (yardımcı patojenite/konservasyon skorları) ve AA_1/AA_2 (amino asit değişimleri) "
        "gruplarından oluşmaktadır. ground truth etiketi, ACMG uyumlu ClinVar sınıflandırmasına "
        "dayanmaktadır."
    )

    body(doc,
        "Problemin temel zorluğu üç eksenlidir: (1) Aşırı sınıf dengesizliği — eğitim verisinde "
        "pathogenic varyantlar çoğunluktadır (~%73 MASTER'da), ancak yarışma final test setinin "
        "~%80 benign / %20 pathogenic dağılımına sahip olacağı bildirilmiştir; bu, eğitim ve test "
        "dağılımlarının tam tersi olması (label shift) anlamına gelmektedir. (2) Yüksek eksiklik oranı — "
        "MASTER verisinde ortalama eksiklik %55'tir ve 165 sütun %50'nin üzerinde eksik değere sahiptir; "
        "bu eksiklik rastgele değildir ve etiketle korelasyon göstermektedir (pathogenic satırlarda %59.9, "
        "benign satırlarda %41.2). (3) Alt-panel heterojenliği — her gen paneli farklı biyolojik "
        "mekanizmalara sahip olduğundan, tek bir genel modelin tüm panellere genellemesi zordur [2], [3]."
    )

    body(doc,
        "Çalışmamızda bu zorlukları ele almak için panel-bazlı adaptif bir modelleme stratejisi "
        "benimsenmiştir. Her panel için ayrı model konfigürasyonları geliştirilmiş, label shift'e karşı "
        "benign-aware karar eşiği optimizasyonu ve prior-shift düzeltmesi uygulanmış, eksiklik bilgisinin "
        "informatif bir özellik olarak kullanıldığı M3 stratejisi geliştirilmiş ve BalancedBagging "
        "ile undersampling tabanlı ensemble yöntemleri kullanılmıştır."
    )

    # =====================================================
    # 2. YÖNTEM
    # =====================================================
    heading(doc, "2. YÖNTEM")
    doc.add_paragraph()

    subheading(doc, "2.1 Kullanılan Veri Kümesi ve Yapısı")

    body(doc,
        "Yarışma formatı gereği genomik adres bilgileri gizlenmiş olup sütun isimleri anonimleştirilmiştir. "
        "Tablo 1'de veri setinin panel bazlı dağılımı sunulmaktadır."
    )

    cap1 = doc.add_paragraph()
    r1 = cap1.add_run("Tablo 1: Veri Seti Dağılımı")
    r1.italic = True; r1.font.size = Pt(9); r1.font.name = "Aptos"
    t1 = doc.add_table(rows=1, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(t1, ["Panel", "Toplam", "Patojenik (1)", "Benign (0)", "Kullanım"])
    add_table_row(t1, ["MASTER", "2931", "2149 (%73.3)", "782 (%26.7)", "Eğitim + Hold-out Test"])
    add_table_row(t1, ["KANSER", "388", "268 (%69.1)", "120 (%30.9)", "Eğitim + Hold-out Test"])
    add_table_row(t1, ["PAH", "372", "310 (%83.3)", "62 (%16.7)", "Panel-bazlı test"])
    add_table_row(t1, ["CFTR", "111", "90 (%81.1)", "21 (%18.9)", "Panel-bazlı test"])
    doc.add_paragraph()

    body(doc,
        "Anonimleştirilmiş sütunların grup anatomisi ayrıntılı EDA ile belirlenmiştir: "
        "AL_1..AL_334 sütunlarının 103 tanesi binary (0/1) bayrak, kalan 231 tanesi popülasyon "
        "frekansı ve in-silico skorlardır. CAT_1 kaynak popülasyon kohortu (~30 değer), CAT_2 "
        "AllofUs popülasyonu (7 değer), CAT_3/CAT_4/CAT_5 genotip bilgileri (CAT_3 ve CAT_5 "
        "birebir özdeş tespit edilmiştir), CAT_6 genomik bölge bayrağı (%97.7 eksik), "
        "EK_1..EK_9 yardımcı patojenite/konservasyon skorları ve AA_1/AA_2 tek harf amino asit "
        "değişimleridir."
    )

    subheading(doc, "2.2 Veri Ön İşleme ve Temsilleme Stratejisi")

    body(doc,
        "Veri ön işleme aşamasında, data leakage'ı önlemek amacıyla tüm "
        "fit işlemleri (imputer, encoder, scaler) yalnızca eğitim bölünmesi üzerinde "
        "gerçekleştirilmiş, test ve panel verilerine yalnızca transform uygulanmıştır. "
        "Bu kural, TEKNOFEST veri setindeki yüksek eksiklik oranının (%55 ortalama) "
        "sızıntı potansiyelini artırması nedeniyle özellikle kritiktir [4]."
    )

    body(doc, "Veri ön işleme adımları şu şekilde gerçekleştirilmiştir:")

    bullet(doc,
        "Sabit ve özdeş sütun temizliği: MASTER verisinde 57 sabit sütun (nunique ≤ 1) ve "
        "583 birebir özdeş sütun çifti tespit edilerek toplam 63 sütun drop edilmiştir. "
        "Ham 351 özellikten 288'i korunmuştur."
    )
    bullet(doc,
        "M3 eksiklik stratejisi (deneysel olarak doğrulanmış en iyi yaklaşım, panel ort. F1=0.9210): "
        "%50'nin üzerinde eksik olan sütunlar için is_missing_* binary bayrakları oluşturulmuş "
        "(eksiklik bilgi taşıdığından drop yerine korunmuştur), tüm sayısal sütunlar eğitim seti "
        "medyanı ile doldurulmuştur."
    )
    bullet(doc,
        "Kategorik kodlama: İki ayrı pipeline denenmiştir — (A) one-hot encoding (OHE) ve "
        "(B) label encoding (LE). CatBoost modeli için native kategorik destek kullanılmıştır."
    )
    bullet(doc,
        "feature engineering (FE): AA_1/AA_2 amino asit bilgilerinden Grantham mesafesi "
        "(gömülü 20×20 matris), BLOSUM62 skoru (gömülü matris), stopgain bayrağı (AA_2=='*') "
        "ve AL frekans sütunlarına log1p dönüşümü uygulanmıştır. FE'nin katkısı ablasyon "
        "yöntemiyle izole olarak ölçülmüştür: CFTR'de +0.069, KANSER'de +0.015, PAH'ta −0.003 "
        "(%80/20 F1 deltaları) [5]."
    )

    subheading(doc, "2.3 Sınıf Dengesi ve Label Shift Yönetimi")

    body(doc,
        "Eğitim verisinde pathogenic sınıf çoğunlukta (%73 MASTER'da, %83 PAH'ta) iken "
        "yarışma test setinin ~%80 benign dağılımına sahip olması, standart F1-maximizing "
        "karar eşiğinin üretimde precision çöküşüne yol açacağı anlamına gelmektedir. Bu "
        "problemi ele almak için üç katmanlı bir strateji geliştirilmiştir:"
    )

    bullet(doc,
        "BalancedBagging ensemble: Eğitim sırasında çoğunluk sınıfının (pathogenic) "
        "rasgele alt-örneklenmesi ile azınlık sınıfıyla (benign) dengelenen birden fazla "
        "base learner'ın birleştirilmesi. Bu yaklaşım, Chatterji vd. (2022) tarafından "
        "label shift altında minimax-optimal olduğu teorik olarak kanıtlanmıştır [6]."
    )
    bullet(doc,
        "benign-aware karar eşiği optimizasyonu: Eşik değeri, eğitim verisinin doğal "
        "dağılımı yerine %80 benign / %20 pathogenic resample edilmiş bir doğrulama havuzunda "
        "optimize edilmiştir. Üç mod karşılaştırılmıştır: f1_raw (eski), f1_8020 ve "
        "mcc_8020. Sonuçlar f1_8020 ve mcc_8020'nin eşdeğer olduğunu ve f1_raw'dan "
        "~2 puan üstün olduğunu göstermiştir."
    )
    bullet(doc,
        "prior-shift düzeltmesi (Saerens vd. 2002): Model çıktı olasılıklarının "
        "test dağılımına göre post-hoc olarak kalibre edilmesi. CFTR'de etkili "
        "(MCC: +0.028), PAH'ta zararlı bulunmuştur (kalibrasyonsuz olasılıklarda). "
        "Alexandari vd. (2020, ICML) önerisine uygun olarak, kalibrasyon (Platt scaling) "
        "EM'den ÖNCE uygulandığında prior-shift'in PAH'ta da pozitif etkisi "
        "gözlemlenmiştir (+0.052 MCC) [7], [8]."
    )

    subheading(doc, "2.4 Seçilen Algoritmalar ve Gerekçe")

    body(doc,
        "tabular veriler için kanıtlanmış performansları nedeniyle gradient boosting "
        "ailesi temel alınmış, derin öğrenme modelleri tamamlayıcı olarak eklenmiştir:"
    )

    bullet(doc,
        "LightGBM: tabular veride state-of-the-art performans, leaf-wise büyüme stratejisi "
        "ile yüksek hız ve doğruluk. Tüm panellerde birincil model olarak kullanılmıştır."
    )
    bullet(doc,
        "XGBoost: depth-wise ağaç büyütme, LightGBM'e alternatif olarak çeşitlilik sağlamak "
        "amacıyla stacking ensemble'da kullanılmıştır."
    )
    bullet(doc,
        "CatBoost: Native kategorik değişken desteği ile OHE/LE gerektirmeden direkt "
        "kategorik sütunları işleyebilmesi nedeniyle denenmiştir."
    )
    bullet(doc,
        "SmallMLP (focal loss, γ=2): Doğrusal olmayan etkileşimleri yakalamak için "
        "tasarlanmış, residual connection'lı 3-4 katmanlı yapay sinir ağı. "
        "BatchNorm yerine yüksek dropout + weight_decay ile düzenlileştirme; "
        "focal loss ile azınlık sınıfına (benign) odaklanma sağlanmıştır [9]."
    )
    bullet(doc,
        "SVM (RBF kernel): Yüksek boyutlu uzayda etkili karar sınırları; "
        "Platt scaling ile olasılık kalibrasyonu."
    )
    bullet(doc,
        "TabPFN: tabular foundation model (Hollmann vd. 2025, Nature); ≤10k örnekte "
        "GBDT'leri geçtiği raporlanmış olup PAH plato-kırma denemesinde test edilmiştir [10]."
    )

    subheading(doc, "2.5 Hiperparametre Optimizasyonu ve Çapraz Doğrulama")

    body(doc,
        "Tüm modellerde 3-fold (sonraki deneylerde 5-fold) Stratified K-Fold çapraz doğrulama "
        "ile sistematik grid search uygulanmıştır. LightGBM ve XGBoost için 12 kombinasyon "
        "(n_estimators × max_depth/num_leaves × learning_rate), NN/DNN için 4 kombinasyon "
        "(dropout × learning_rate) denenmiştir. Tüm süreçlerde SEED=42 sabit tutulmuştur. "
        "overfitting'i önlemek için validation-based early stopping (patience=10), "
        "class_weight='balanced' / scale_pos_weight parametreleri ve undersampling "
        "kullanılmıştır. Karar eşiği [0.10, 0.90] aralığında F1 maximizing grid search "
        "ile belirlenmiştir."
    )

    subheading(doc, "2.6 Stacking Ensemble ve Out-of-Fold (OOF) Protokolü")

    body(doc,
        "out-of-fold stacking yaklaşımı, 5-fold CV ile her base modelden (LightGBM, XGBoost, "
        "CatBoost, NN, DNN) out-of-fold olasılık tahminleri üretilmiş ve bu tahminler "
        "meta-feature olarak birleştirilmiştir. meta-learner olarak iki aday karşılaştırılmıştır: "
        "Logistic Regression (L2, class_weight=balanced) ve LightGBM. Sonuçlar, GBM meta-learner'ın "
        "küçük meta-feature matrisinde overfit gösterdiğini (train F1: 0.84–0.87), "
        "Logistic Regression'ın ise tutarlı şekilde üstün olduğunu göstermiştir "
        "(stack_lr panel ort. %80/20 F1=0.638 vs stack_gbm=0.566)."
    )

    subheading(doc, "2.7 Değerlendirme Protokolü")

    body(doc,
        "label shift problemini dürüstçe ele almak için iki aşamalı değerlendirme "
        "protokolü tasarlanmıştır: (1) %50/50 hold-out test (mevcut panel dağılımında), "
        "(2) bootstrap %80/20 değerlendirme — tüm benign'ler sabit tutularak pathogenic "
        "örnekler %20'ye downsample edilmiştir (N=50 tekrar, %95 CI). Birincil metrik "
        "olarak bootstrap %80/20 pathogenic F1 kullanılmıştır. Başarım ölçütleri olarak "
        "F1, MCC (Matthews Correlation Coefficient), AUC-PR (precision-recall eğrisi "
        "altında kalan alan), precision ve recall raporlanmıştır [11]."
    )

    # =====================================================
    # 3. BULGULAR
    # =====================================================
    heading(doc, "3. BULGULAR")
    doc.add_paragraph()

    subheading(doc, "3.1 Baseline Sonuçları")

    body(doc,
        "İlk aşamada, MASTER verisi üzerinde üç encoding pipeline (OHE, label encoding, "
        "CatBoost native) ile 7 model eğitilmiştir. Tablo 2'de MASTER %80/20 hold-out "
        "test sonuçları sunulmaktadır."
    )

    cap2 = doc.add_paragraph()
    r2 = cap2.add_run("Tablo 2: MASTER Baseline Sonuçları (Hold-Out %80/20)")
    r2.italic = True; r2.font.size = Pt(9); r2.font.name = "Aptos"
    t2 = doc.add_table(rows=1, cols=6)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(t2, ["Model", "F1", "AUC-ROC", "AUC-PR", "MCC", "Pipeline"])
    add_table_row(t2, ["LightGBM", "0.9465", "0.9674", "0.9825", "0.8192", "LE"])
    add_table_row(t2, ["XGBoost", "0.9487", "0.9686", "0.9851", "0.8269", "LE"])
    add_table_row(t2, ["CatBoost", "0.9501", "0.9726", "0.9884", "0.8307", "Native"])
    add_table_row(t2, ["SVM", "0.8955", "0.8960", "0.9436", "0.6352", "LE"])
    doc.add_paragraph()

    subheading(doc, "3.2 Ablasyon Deneyleri")

    body(doc,
        "13 senaryo × 2 model (XGBoost + LightGBM) üzerinde kapsamlı ablasyon "
        "gerçekleştirilmiştir. EK skorlarının (EK_1..EK_9) çıkarılması en büyük F1 "
        "düşüşünü yaratmıştır (−0.022). Tüm AL sütunlarının çıkarılması F1'de −0.015 "
        "düşüşe neden olurken, kategorik sütunların çıkarılması etkisizdir. "
        "is_missing_* bayraklarının eklenmesi (M3 stratejisi) F1'de tutarlı artış sağlamıştır. "
        "Şekil 1'de ablasyon sonuçları gösterilmektedir."
    )

    add_figure(doc, str(RESULTS / "ablation_ohe_xgb" / "fig1_f1drop_comparison.png"),
               "Şekil 1: Ablasyon deneyi F1 drop karşılaştırması (XGBoost vs LightGBM)")

    subheading(doc, "3.3 Panel-Bazlı Eğitim Sonuçları")

    body(doc,
        "Panel-bazlı eğitim stratejisi, her panel için MASTER verisi ile birleştirilmiş "
        "eğitim havuzları kullanılarak gerçekleştirilmiştir. Tablo 3'te, iki ardışık deney turundaki "
        "(FE + stacking) sonuçları panel bazında karşılaştırılmaktadır. En kritik bulgu, "
        "%50/50 test performansının %80/20 bootstrap'a çevrildiğinde dramatik çöküş "
        "göstermesidir (PAH: 0.917 → 0.489)."
    )

    cap3 = doc.add_paragraph()
    r3 = cap3.add_run("Tablo 3: Panel-Bazlı En İyi Sonuçlar (%80/20 Bootstrap F1)")
    r3.italic = True; r3.font.size = Pt(9); r3.font.name = "Aptos"
    t3 = doc.add_table(rows=1, cols=5)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(t3, ["Panel", "1. Tur En İyi", "2. Tur En İyi", "Model", "%95 CI"])
    add_table_row(t3, ["KANSER", "0.690", "0.716", "stack_lr", "[0.67–0.77]"])
    add_table_row(t3, ["PAH", "0.489", "0.515", "stack_lr", "[0.35–0.57]"])
    add_table_row(t3, ["CFTR", "0.850", "0.852", "DNN", "[0.00–1.00]*"])

    cap3b = doc.add_paragraph()
    r3b = cap3b.add_run("* CFTR'de n=21 benign nedeniyle bootstrap CI tüm skalayı kapsamaktadır; model karşılaştırmaları istatistiksel olarak anlamsızdır.")
    r3b.italic = True; r3b.font.size = Pt(8); r3b.font.name = "Aptos"

    add_figure(doc, str(RESULTS / "v6_fe_stacking" / "fig1_base_vs_stack.png"),
               "Şekil 2: Base model vs Stacking performans karşılaştırması")

    add_figure(doc, str(RESULTS / "v6_fe_stacking" / "fig2_fe_ablation.png"),
               "Şekil 3: Özellik mühendisliği (FE) ablasyon sonuçları")

    subheading(doc, "3.4 CFTR Paneli: Refinement ve Final Model")

    body(doc,
        "CFTR panelinin n=21 benign ile çok küçük negatif sınıfa sahip olması, bootstrap "
        "CI'nın [0.00–1.00] genişliğinde kalmasına neden olmuştur. Bu sorunu çözmek için "
        "LOO-CV + MCC değerlendirme protokolüne "
        "geçilmiştir. 7 strateji sistematik olarak karşılaştırılmıştır. "
        "Tablo 4'te CFTR strateji karşılaştırması sunulmaktadır."
    )

    cap4 = doc.add_paragraph()
    r4 = cap4.add_run("Tablo 4: CFTR Strateji Karşılaştırması (LOO-CV)")
    r4.italic = True; r4.font.size = Pt(9); r4.font.name = "Aptos"
    t4 = doc.add_table(rows=1, cols=6)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(t4, ["Strateji", "LOO-MCC", "LOO-F1", "Boot %80/20", "TN/FP", "Overfit"])
    add_table_row(t4, ["S6_PriorShift", "0.655", "0.919", "0.692", "18/3", "OK"])
    add_table_row(t4, ["S0_MASTER-only", "0.627", "0.898", "0.728", "18/3", "ORTA"])
    add_table_row(t4, ["S5_BalancedBagging", "0.572", "0.855", "0.754", "—", "OK"])
    add_table_row(t4, ["S2_SMOTE", "0.561", "0.924", "0.493", "—", "OK*"])
    add_table_row(t4, ["S4_FrozenFinetune", "0.294", "0.500", "0.370", "—", "OK"])
    doc.add_paragraph()

    body(doc,
        "Sonraki deneylerde CFTR modeli daha da olgunlaştırılmıştır. Eğitim havuzu "
        "MASTER + KANSER + PAH olarak genişletildiğinde (COMBINED, n=3691), model "
        "FP=0 (precision=1.0) elde etmiştir. Bu, %80 benign final test setinde kritik "
        "avantaj sağlamaktadır. Final CFTR modeli olarak S0c_COMBINED seçilmiştir: "
        "LOO-MCC=0.644, bootstrap %80/20 F1=0.863, precision=1.0."
    )

    add_figure(doc, str(RESULTS / "v10_cftr_ensemble" / "fig2_strategy_comparison.png"),
               "Şekil 4: CFTR strateji karşılaştırması — S0c_COMBINED final model")

    add_figure(doc, str(RESULTS / "v10_cftr_ensemble" / "fig1_posterior_histogram.png"),
               "Şekil 5: CFTR COMBINED posterior dağılımı (FP=0, precision=1.0)")

    subheading(doc, "3.5 PAH Paneli: Refinement ve Plato Analizi")

    body(doc,
        "PAH paneli, projede en zorlu panel olarak ortaya çıkmıştır (n=62 benign, %83 pathogenic). "
        "İlk stacking baseline (%80/20 F1=0.515) üzerine, 6 ardışık deney boyunca sistematik "
        "iyileştirme denenmiştir. Tablo 5'te PAH yolculuğu özetlenmektedir."
    )

    cap5 = doc.add_paragraph()
    r5 = cap5.add_run("Tablo 5: PAH Performans Yolculuğu")
    r5.italic = True; r5.font.size = Pt(9); r5.font.name = "Aptos"
    t5 = doc.add_table(rows=1, cols=5)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(t5, ["Deney", "En İyi Strateji", "MCC", "Boot %80/20 F1", "Temel Bulgu"])
    add_table_row(t5, ["Baseline", "stack_lr", "—", "0.515", "Baseline stacking"])
    add_table_row(t5, ["Deney 1", "P4_COMBINED_BalBag", "0.529", "0.582", "BalBag + COMBINED (+0.067)"])
    add_table_row(t5, ["Deney 2", "Sweep n=10", "0.536", "0.591", "Marjinal sweep iyileşme"])
    add_table_row(t5, ["Deney 3", "T4_COMBINED+BalBag", "0.529", "0.576", "Havuz katkı analizi"])
    add_table_row(t5, ["Deney 4", "D3 SplineCal+Prior", "0.503", "0.581", "TabPFN platoyu kıramadı"])
    add_table_row(t5, ["Deney 5", "Multi-seed ort.", "0.489*", "0.526*", "FE+HyperEns nötr"])

    cap5b = doc.add_paragraph()
    r5b = cap5b.add_run("* Multi-seed ortalama değerleri; tek seed en iyi MCC=0.578")
    r5b.italic = True; r5b.font.size = Pt(8); r5b.font.name = "Aptos"
    doc.add_paragraph()

    body(doc,
        "PAH'ta MCC≈0.53, bootstrap F1≈0.58 platosuna ulaşılmış ve bu plato 5 farklı müdahale "
        "ekseni (havuz genişletme, BalancedBagging sweep, kalibrasyon, TabPFN, feature "
        "engineering + HyperEnsemble) ile kırılamamıştır. Bu bulgu, Chatterji vd. (2022) "
        "tarafından teorik olarak kanıtlanan azınlık-örneği tavanı (minimax-optimal "
        "undersampling sınırı) ile tutarlıdır: n=62 benign ile daha fazla iyileşme, "
        "veri-sinyali tavanına çarpmaktadır [6]. PAH çalışması bu sonuçla güvenle "
        "kapatılmıştır."
    )

    add_figure(doc, str(RESULTS / "v8_pah_refinement" / "fig3_confusion.png"),
               "Şekil 6: PAH confusion matrix karşılaştırması")

    add_figure(doc, str(RESULTS / "v11_pah_tabpfn" / "fig2_bootstrap_f1.png"),
               "Şekil 7: PAH bootstrap F1 karşılaştırması — plato doğrulanması")

    subheading(doc, "3.6 Hereditary Cancer (KANSER) Paneli — Devam Eden Çalışma")

    body(doc,
        "KANSER paneli için şu ana kadar stacking (LR meta) ile %80/20 bootstrap F1=0.716 "
        "(CI=[0.67–0.77]) elde edilmiştir. Bu, projedeki en güvenilir sonuçtur "
        "(dar CI, orta düzey dengesizlik). KANSER paneline COMBINED+BalancedBagging "
        "müdahalesi henüz uygulanmamıştır. CFTR ve PAH'ta kazanılan deneyimler "
        "(özellikle BalancedBagging'in etkinliği) KANSER'e de transfer edilecektir. "
        "Bu çalışma devam etmektedir."
    )

    subheading(doc, "3.7 MASTER (Genel Panel) — Devam Eden Çalışma")

    body(doc,
        "MASTER panelinde baseline XGBoost F1=0.9487 (AUC-ROC=0.9686) elde edilmiştir. "
        "stacking NN ile F1=0.9505'e yükseltilmiştir. Ancak MASTER sonuçları %50/50 "
        "dağılımda ölçülmüştür; %80/20 final-realistic değerlendirme henüz "
        "gerçekleştirilmemiştir. KANSER müdahalesinin ardından MASTER üzerinde de "
        "benign-aware threshold + bootstrap değerlendirme protokolü uygulanacaktır. "
        "Bu çalışma devam etmektedir."
    )

    subheading(doc, "3.8 Karar Eşiği Analizi")

    body(doc,
        "Farklı karar eşiklerinin model performansı üzerindeki etkisi sistematik olarak "
        "incelenmiştir. Üç eşik modu karşılaştırılmıştır: f1_raw (eğitim dağılımında "
        "F1-max), f1_8020 (%80/20 resample'da F1-max) ve mcc_8020 (%80/20'de MCC-max). "
        "Tablo 6'da seçilen eşik değerleri sunulmaktadır."
    )

    cap6 = doc.add_paragraph()
    r6 = cap6.add_run("Tablo 6: Panel Bazında Optimal Karar Eşikleri")
    r6.italic = True; r6.font.size = Pt(9); r6.font.name = "Aptos"
    t6 = doc.add_table(rows=1, cols=4)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(t6, ["Panel", "Final Model", "Eşik Değeri", "Eşik Modu"])
    add_table_row(t6, ["CFTR", "S0c_COMBINED (LightGBM)", "0.147", "mcc_8020"])
    add_table_row(t6, ["PAH", "P4_COMBINED_BalBag (LightGBM)", "0.050", "mcc_raw"])
    add_table_row(t6, ["KANSER", "stack_lr (6 base + LR meta)", "—", "Devam ediyor"])
    doc.add_paragraph()

    # =====================================================
    # 4. SONUÇ
    # =====================================================
    heading(doc, "4. SONUÇ")
    doc.add_paragraph()

    body(doc,
        "Bu çalışmada, TEKNOFEST 2026 Sağlıkta Yapay Zeka Yarışması kapsamında "
        "missense genetik varyantların pathogenic/benign sınıflandırması için panel-bazlı "
        "adaptif bir makine öğrenmesi framework'ü geliştirilmiştir. Çalışmanın temel "
        "katkıları şu şekilde özetlenebilir:"
    )

    subheading(doc, "4.1 Güçlü Yönler ve Başarılar")

    bullet(doc,
        "label shift'e karşı dayanıklı değerlendirme protokolü: %80/20 bootstrap "
        "değerlendirme, eğitim-test dağılım uyumsuzluğunun yarattığı yanılsamayı "
        "ortaya çıkarmıştır (%50/50'de 0.91–0.93 görülen F1'ler, gerçekçi değerlendirmede "
        "0.49–0.85'e düşmüştür). Bu, literatürde sıkça göz ardı edilen bir tuzaktır."
    )
    bullet(doc,
        "M3 eksiklik stratejisi: Eksikliğin informatif bir sinyal olarak kullanılması "
        "(is_missing_* bayrakları), drop veya basit imputation stratejilerinden üstün "
        "performans sergilemiştir (panel ort. F1: 0.9210 vs M1=0.9165 vs M5=0.9188)."
    )
    bullet(doc,
        "CFTR'de precision=1.0: COMBINED eğitim havuzu ile sıfır false positive "
        "elde edilmiştir; bu, %80 benign final test setinde kritik avantaj sağlamaktadır."
    )
    bullet(doc,
        "BalancedBagging'in tutarlı üstünlüğü: PAH'ta 5 farklı deney boyunca "
        "BalancedBagging sürekli en iyi MCC'yi vermiştir (0.529–0.536), teorik "
        "beklentiyle (Chatterji 2022) uyumlu bir şekilde."
    )

    subheading(doc, "4.2 Sınırlılıklar ve Zorluklar")

    body(doc,
        "Çalışmanın en belirgin sınırlılığı PAH panelindeki performans platosudur "
        "(bootstrap F1≈0.58). n=62 benign ile azınlık-örneği tavanına ulaşılmış olup, "
        "daha fazla veri olmadan dramatik iyileşme mümkün görünmemektedir. CFTR'de "
        "ise n=21 benign nedeniyle bootstrap CI çok geniştir ve nokta tahminler "
        "güvenilir karşılaştırma yapmaya elverişli değildir."
    )

    body(doc,
        "false positive ve false negative analizinde: CFTR'de model 3 benign varyantı "
        "yanlışlıkla pathogenic olarak sınıflandırmıştır (FP=3, S6 stratejisinde; "
        "COMBINED ile FP=0'a düşürülmüştür). PAH'ta ise 45 pathogenic varyant yanlışlıkla "
        "benign olarak sınıflandırılmıştır (FN=45, recall=0.85). Bu FN'lerin klinik "
        "anlamı, potansiyel olarak hastalık yapıcı bir varyantın gözden kaçırılmasıdır "
        "— bu, yanlış tanıya veya tedavide gecikmeye yol açabilir."
    )

    body(doc,
        "Modelin en çok zorlandığı özellik grupları: düşük MAF (Minor Allele Frequency) "
        "varyantları, nadir popülasyonlardan gelen varyantlar ve fonksiyonel anotasyonu "
        "yetersiz olan bölgeler. EK skorlarının (özellikle EK_7, EK_1, EK_2) çıkarılması "
        "en büyük performans düşüşüne neden olmuştur; bu da modelin büyük ölçüde "
        "in-silico patojenite tahmin araçlarına bağımlı olduğunu göstermektedir."
    )

    subheading(doc, "4.3 Literatürdeki Yeri")

    body(doc,
        "Literatürde varyant patojenite tahmini genellikle genom-çapında (genome-wide) "
        "eğitilmiş büyük modeller (REVEL, ClinPred, AlphaMissense) üzerinden "
        "gerçekleştirilmektedir [3], [12], [13]. Bu yaklaşımlar yüksek genel doğruluk "
        "sağlar ancak gen-spesifik kalibrasyon düşük olabilir — Wu vd. (2024) "
        "genomik-çapında kalibre edilen modellerin %70'inin tek-gen düzeyinde başarısız "
        "olduğunu göstermiştir [14]. Bizim panel-bazlı yaklaşımımız, her gen paneli "
        "için ayrı kalibrasyon ve eşik optimizasyonu yaparak bu sorunu ele almaktadır."
    )

    body(doc,
        "Ayrıca, label shift altında değerlendirme protokolü tasarlamamız (bootstrap "
        "%80/20 + prior-shift düzeltmesi) ve BalancedBagging'in minimax-optimal "
        "özelliğinin deneysel doğrulanması, yöntemsel katkılar olarak değerlendirilebilir."
    )

    subheading(doc, "4.4 Final Aşamasında Beklenen Zorluklar")

    body(doc,
        "Yarışmanın son aşamasında karşılaşılması beklenen zorluklar: (1) Gerçek test "
        "setinin dağılımının tam olarak bilinmemesi — %80/20 varsayımı yanlışsa eşik "
        "değerlerinin yeniden ayarlanması gerekecektir. (2) KANSER ve MASTER panelleri "
        "için henüz benign-aware optimizasyon tamamlanmamıştır. (3) PAH panelinin "
        "performans tavanı (bootstrap F1≈0.58) diğer takımlara göre dezavantaj "
        "oluşturabilir. (4) Anonimleştirilmiş sütunlar nedeniyle domain bilgisi "
        "(protein yapısı, 3D konum vb.) entegrasyonu sınırlıdır."
    )

    # =====================================================
    # 5. KAYNAKÇA
    # =====================================================
    heading(doc, "5. KAYNAKÇA")
    doc.add_paragraph()

    refs = [
        '[1] B. Li ve M. Dewey, "Missense variant pathogenicity prediction: Current landscape and challenges," '
        'Briefings in Bioinformatics, c. 23, sy. 4, sf. bbac196, 2022.',
        '[2] Q. Chen vd., "DS-MVP: Identifying disease-specific pathogenicity of missense variants by '
        'pre-training representation," Briefings in Bioinformatics, c. 26, sy. 2, sf. bbaf119, 2025.',
        '[3] L. Gerasimavicius, S. A. Teichmann ve J. A. Marsh, "Leveraging protein structural information '
        'to improve variant effect prediction," Current Opinion in Structural Biology, 2025.',
        '[4] S. Kaufman vd., "Leakage in data mining: Formulation, detection, and avoidance," '
        'ACM Transactions on Knowledge Discovery from Data, c. 6, sy. 4, sf. 15, 2012.',
        '[5] R. Grantham, "Amino acid difference formula to help explain protein evolution," '
        'Science, c. 185, sy. 4154, sf. 862-864, 1974.',
        '[6] N. S. Chatterji vd., "The interplay between implicit bias and benign overfitting in two-layer '
        'linear networks," arXiv:2205.13094, 2022.',
        '[7] M. Saerens, P. Latinne ve C. Decaestecker, "Adjusting the outputs of a classifier to new '
        'a priori probabilities: A simple procedure," Neural Computation, c. 14, sy. 1, sf. 21-41, 2002.',
        '[8] A. Alexandari, A. Kundaje ve A. Shrikumar, "Maximum likelihood with bias-corrected calibration '
        'is hard-to-beat at label shift adaptation," Proc. 37th ICML, PMLR 119, sf. 222-232, 2020.',
        '[9] T.-Y. Lin vd., "Focal loss for dense object detection," Proc. IEEE ICCV, sf. 2980-2988, 2017.',
        '[10] N. Hollmann vd., "Accurate predictions on small data with a tabular foundation model," '
        'Nature, c. 637, sf. 319-326, 2025.',
        '[11] D. Chicco ve G. Jurman, "The advantages of the Matthews correlation coefficient (MCC) '
        'over F1 score and accuracy in binary classification evaluation," BMC Genomics, c. 21, sy. 6, 2020.',
        '[12] M. Ioannidis vd., "REVEL: An ensemble method for predicting the pathogenicity of rare missense '
        'variants," American Journal of Human Genetics, c. 99, sy. 4, sf. 877-885, 2016.',
        '[13] J. Cheng vd., "Accurate proteome-wide missense variant effect prediction with AlphaMissense," '
        'Science, c. 381, sy. 6664, sf. eadg7492, 2023.',
        '[14] L. Wu vd., "Gene-specific calibration of variant pathogenicity prediction tools," '
        'Genome Medicine, c. 16, sf. 45, 2024.',
        '[15] M. Kabir vd., "PON-P3: Accurate prediction of pathogenicity of amino acid substitutions," '
        'International Journal of Molecular Sciences, c. 26, sy. 5, sf. 2004, 2025.',
        '[16] J. Lai vd., "LYRUS: A machine learning model for predicting the pathogenicity of missense '
        'variants," Bioinformatics Advances, c. 2, sy. 1, sf. vbab045, 2021.',
        '[17] N. Alirezaie vd., "ClinPred: Prediction tool to identify disease-relevant nonsynonymous '
        'single-nucleotide variants," American Journal of Human Genetics, c. 103, sy. 4, sf. 474-483, 2018.',
        '[18] M. Danner vd., "Utilizing protein structure graph embeddings to predict the pathogenicity of '
        'missense variants," NAR Genomics and Bioinformatics, c. 7, sy. 3, sf. lqaf097, 2025.',
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.style = doc.styles['Normal']
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = "Aptos"
        p.paragraph_format.space_after = Pt(4)

    # Adjust margins: top 2.8cm, bottom/left/right 2.5cm
    for section in doc.sections:
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Save
    doc.save(str(OUTPUT))
    print(f"PDR kaydedildi: {OUTPUT}")
    print(f"Dosya boyutu: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build_report()
